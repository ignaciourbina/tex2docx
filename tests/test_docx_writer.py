"""Tests for DOCX writer (IR -> .docx)."""

import os

from docx import Document as DocxDocument

from tex2docx.ir import (
    Abstract, Comment, Document, Formatted, FormatKind,
    Image, List, ListItem, ListKind, Math, Metadata,
    Paragraph, Preamble, RawLatex, Reference, RefKind,
    Section, SectionLevel, Table, TableRow, TableCell, Text,
)
from tex2docx.docx.writer import DocxWriter


def _write_and_read(ir_doc, tmp_path):
    """Write IR to DOCX, then read it back with python-docx."""
    out = str(tmp_path / "test.docx")
    DocxWriter().write(ir_doc, out)
    assert os.path.isfile(out)
    return DocxDocument(out)


def test_simple_paragraph(tmp_path):
    ir = Document(children=[
        Paragraph(children=[Text(content="Hello, world!")])
    ])
    doc = _write_and_read(ir, tmp_path)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "Hello, world!" in texts


def test_metadata(tmp_path):
    ir = Document(children=[
        Metadata(title="My Title", author="Author Name", date="2026"),
    ])
    doc = _write_and_read(ir, tmp_path)
    assert doc.core_properties.title == "My Title"
    assert doc.core_properties.author == "Author Name"


def test_sections(tmp_path):
    ir = Document(children=[
        Section(level=SectionLevel.SECTION, title="Introduction"),
        Section(level=SectionLevel.SUBSECTION, title="Background"),
    ])
    doc = _write_and_read(ir, tmp_path)
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(headings) == 2
    assert headings[0].text == "Introduction"
    assert headings[1].text == "Background"


def test_bold_italic(tmp_path):
    ir = Document(children=[
        Paragraph(children=[
            Text(content="Normal "),
            Formatted(kind=FormatKind.BOLD, children=[Text(content="bold")]),
            Text(content=" and "),
            Formatted(kind=FormatKind.ITALIC, children=[Text(content="italic")]),
        ])
    ])
    doc = _write_and_read(ir, tmp_path)
    para = [p for p in doc.paragraphs if p.text.strip()][0]
    runs = para.runs
    bold_runs = [r for r in runs if r.bold]
    italic_runs = [r for r in runs if r.italic]
    assert any("bold" in r.text for r in bold_runs)
    assert any("italic" in r.text for r in italic_runs)


def test_bullet_list(tmp_path):
    ir = Document(children=[
        List(kind=ListKind.BULLETED, children=[
            ListItem(children=[Text(content="First")]),
            ListItem(children=[Text(content="Second")]),
        ])
    ])
    doc = _write_and_read(ir, tmp_path)
    bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(bullet_paras) == 2


def test_numbered_list(tmp_path):
    ir = Document(children=[
        List(kind=ListKind.NUMBERED, children=[
            ListItem(children=[Text(content="One")]),
            ListItem(children=[Text(content="Two")]),
        ])
    ])
    doc = _write_and_read(ir, tmp_path)
    num_paras = [p for p in doc.paragraphs if p.style.name == "List Number"]
    assert len(num_paras) == 2


def test_table(tmp_path):
    ir = Document(children=[
        Table(col_spec="|c|c|", children=[
            TableRow(children=[
                TableCell(children=[Text(content="A")]),
                TableCell(children=[Text(content="B")]),
            ]),
            TableRow(children=[
                TableCell(children=[Text(content="1")]),
                TableCell(children=[Text(content="2")]),
            ]),
        ])
    ])
    doc = _write_and_read(ir, tmp_path)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 2
    assert doc.tables[0].cell(0, 0).text == "A"
    assert doc.tables[0].cell(1, 1).text == "2"


def test_reference_markers(tmp_path):
    ir = Document(children=[
        Paragraph(children=[
            Text(content="See "),
            Reference(kind=RefKind.CITE, key="paper2024"),
        ])
    ])
    doc = _write_and_read(ir, tmp_path)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "[cite:paper2024]" in full_text


def test_comment_hidden(tmp_path):
    ir = Document(children=[
        Comment(content="This is a comment"),
    ])
    doc = _write_and_read(ir, tmp_path)
    comment_paras = [p for p in doc.paragraphs
                     if p.style and p.style.name == "LaTeX Comment"]
    assert len(comment_paras) == 1


def test_raw_latex_preserved(tmp_path):
    ir = Document(children=[
        RawLatex(content="\\newcommand{\\foo}{bar}"),
    ])
    doc = _write_and_read(ir, tmp_path)
    raw_paras = [p for p in doc.paragraphs
                 if p.style and p.style.name == "LaTeX Raw"]
    assert len(raw_paras) == 1


def test_preamble_embedded(tmp_path):
    ir = Document(children=[
        Preamble(
            document_class="article",
            class_options=["12pt"],
            packages=[("graphicx", []), ("amsmath", [])],
        ),
        Paragraph(children=[Text(content="Body text")]),
    ])
    doc = _write_and_read(ir, tmp_path)
    keywords = doc.core_properties.keywords or ""
    assert "doctex_preamble:" in keywords
    assert "article" in keywords
    assert "graphicx" in keywords
